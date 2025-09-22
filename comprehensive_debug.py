# -*- coding: utf-8 -*-
"""
Comprehensive debug script to analyze Word document structure and find sources of Arabic time formats
"""
import sys
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Need to install python-docx library: pip install python-docx", file=sys.stderr)
    sys.exit(1)

def analyze_document_structure(doc_path):
    """Analyze all content in the Word document to find Arabic time patterns"""
    print(f"\n=== ANALYZING DOCUMENT: {doc_path} ===")

    if not Path(doc_path).exists():
        print(f"Document not found: {doc_path}")
        return

    doc = Document(doc_path)

    # Arabic time patterns to look for
    arabic_patterns = [
        'من', 'إلى', 'ىلإ', 'الى', 'سا',
        'من07:00إلى08:00', 'نم07:00ىلإ08:00',
        '15الى16سا', '07:00', '08:00'
    ]

    print(f"\nDocument has {len(doc.tables)} tables and {len(doc.paragraphs)} paragraphs")

    # Check paragraphs
    print(f"\n--- PARAGRAPHS ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            has_arabic_pattern = any(pattern in text for pattern in arabic_patterns)
            if has_arabic_pattern or i < 5:  # Show first 5 and any with Arabic patterns
                print(f"Para {i}: '{text}' {'⚠️' if has_arabic_pattern else ''}")

    # Check tables in detail
    print(f"\n--- TABLES ---")
    for table_idx, table in enumerate(doc.tables):
        print(f"\nTable {table_idx}: {len(table.rows)} rows × {len(table.columns)} columns")

        # Sample first few rows and look for patterns
        for row_idx, row in enumerate(table.rows):
            if row_idx < 3 or row_idx >= len(table.rows) - 3:  # First 3 and last 3 rows
                print(f"  Row {row_idx}:")
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        has_arabic_pattern = any(pattern in text for pattern in arabic_patterns)
                        print(f"    Cell {cell_idx}: '{text}' {'⚠️' if has_arabic_pattern else ''}")

            # For tables 2 and 3, check all hour cells (column 2)
            if table_idx in [2, 3] and row_idx >= 2:  # Skip headers
                if len(row.cells) > 2:
                    hour_cell = row.cells[2]
                    hour_text = hour_cell.text.strip()
                    if hour_text:
                        has_arabic_pattern = any(pattern in hour_text for pattern in arabic_patterns)
                        if has_arabic_pattern:
                            print(f"  ⚠️ Row {row_idx}, Hour Cell: '{hour_text}'")

    # Check XML content for hidden patterns
    print(f"\n--- XML CONTENT ANALYSIS ---")
    try:
        xml_content = str(doc._element.xml)
        found_patterns = []
        for pattern in arabic_patterns:
            if pattern in xml_content:
                count = xml_content.count(pattern)
                found_patterns.append(f"'{pattern}': {count} times")

        if found_patterns:
            print("Found Arabic patterns in XML:")
            for pattern_info in found_patterns:
                print(f"  {pattern_info}")
        else:
            print("No Arabic time patterns found in XML")
    except Exception as e:
        print(f"Could not analyze XML content: {e}")

    # Check runs and formatting
    print(f"\n--- DETAILED RUN ANALYSIS ---")
    for table_idx, table in enumerate(doc.tables):
        if table_idx in [2, 3]:  # Focus on tables 2 and 3
            print(f"\nTable {table_idx} detailed analysis:")
            for row_idx, row in enumerate(table.rows[2:5], 2):  # Check first few data rows
                if len(row.cells) > 2:
                    hour_cell = row.cells[2]
                    print(f"  Row {row_idx}, Hour Cell:")
                    for para_idx, para in enumerate(hour_cell.paragraphs):
                        print(f"    Para {para_idx}: '{para.text}'")
                        for run_idx, run in enumerate(para.runs):
                            print(f"      Run {run_idx}: '{run.text}'")

if __name__ == "__main__":
    # Analyze the latest corrected document
    analyze_document_structure("test_corrected_final.docx")
