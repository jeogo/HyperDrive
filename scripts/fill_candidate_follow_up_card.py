  # -*- coding: utf-8 -*-
"""
ملء استمارة "بطاقة المتابعة للمترشح" باستبدال الحقول النصية
وإدراج تواريخ وساعات في جداول الدروس.

البنية المكتشفة:
- Table 0: معلومات شخصية (7 صف × 3 عمود) - يحتوي placeholders
- Table 1: جدول الامتحانات (17 صف × 5 عمود) - يبقى فارغ
- Table 2: دروس القانون 30 تاريخ فقط (theory lessons - dates only) - يملأ دائماً
- Table 3: دروس التطبيق 30 تاريخ + ساعة (practical lessons - dates + hours) - يملأ فقط إذا اجتاز الطالب اختبار قانون المرور

منطق التواريخ (محدث):
- Table 2: 30 تاريخ فقط دروس قانون المرور (يبدأ من تاريخ البداية) - يملأ دائماً
- Table 3: 30 تاريخ + ساعة دروس التطبيق (يبدأ بعد 7 أيام من آخر تاريخ في Table 2) - يملأ فقط إذا نجح في اختبار القانون

منطق الساعات (محدث حسب طلب المستخدم):
- منطق تعيين الساعات حسب التاريخ: أول مترشح في تاريخ معين يحصل على 07-08
- ثاني مترشح في نفس التاريخ: 08-09، ثالث مترشح: 09-10... وهكذا
- تواريخ مختلفة: إعادة تشغيل من 07-08 لأول مترشح في التاريخ الجديد
- يتم حفظ الحجوزات في ملف schedule_reservations.json منظم حسب التاريخ
- لا يتم تسجيل مترشحين اثنين في نفس الساعة ونفس التاريخ
- إذا امتلأت جميع الساعات في يوم معين (10 مترشحين) يتم وضع علامة 'FULL'

المتطلبات الأخرى:
- Category دائماً = 'B'
- Table 2: 30 تاريخ فقط (دروس قانون المرور) - يملأ دائماً
- Table 3: 30 تاريخ + ساعات (دروس التطبيق مع فجوة 7 أيام) - يملأ فقط إذا اجتاز اختبار القانون
- تجنب الجمعة والسبت
- ساعات ضمن نافذة 07 إلى 17 (آخر فتحة 16-17)
- كل مترشح له نفس الساعة في كل حصصه (ما لم تنفذ كل الفتحات)
- تنسيق الساعات: XX-XX (مثل 07-08، 08-09) بدون النص العربي

منطق التحقق من اجتياز اختبار قانون المرور:
- يتم فحص البيانات الواردة للتحقق من حالة اختبار قانون المرور (tests.trafficLawTest.passed)
- إذا كان الطالب قد اجتاز الاختبار: يتم ملء الجدول النظري (تواريخ فقط) + الجدول العملي (تواريخ + ساعات)
- إذا لم يجتز الطالب الاختبار بعد: يتم ملء الجدول النظري (تواريخ فقط) وترك الجدول العملي فارغاً
"""
from __future__ import annotations
import argparse
import json
import sys
import re
from pathlib import Path
from datetime import datetime, timedelta
from typing import Dict, Any, List, Tuple

def safe_print(*args, **kwargs):
    """Safe print function that handles Unicode encoding issues on Windows"""
    try:
        print(*args, **kwargs)
    except UnicodeEncodeError:
        # Fallback to ASCII-safe output
        safe_args = []
        for arg in args:
            if isinstance(arg, str):
                safe_args.append(arg.encode('ascii', 'ignore').decode('ascii'))
            else:
                safe_args.append(str(arg))
        print(*safe_args, **kwargs)

try:
    from docx import Document
except ImportError:
    print("[ERROR] Need to install python-docx library: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    from docx2pdf import convert
except ImportError:
    convert = None

# Schedule reservations management
RESERVATIONS_FILE = Path('schedule_reservations.json')

def load_reservations():
    """Load existing hour reservations from JSON file"""
    if RESERVATIONS_FILE.exists():
        try:
            with open(RESERVATIONS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except:
            return {}
    return {}

def save_reservations(reservations):
    """Save hour reservations to JSON file with format validation"""
    try:
        # Validate and clean all hour formats before saving
        cleaned_reservations = clean_reservation_formats(reservations)
        with open(RESERVATIONS_FILE, 'w', encoding='utf-8') as f:
            json.dump(cleaned_reservations, f, indent=2, ensure_ascii=False)
    except Exception as e:
        print(f"[WARN] Failed to save reservations: {e}")

def clean_reservation_formats(reservations):
    """Clean all hour formats in reservation data to ensure consistent XX-XX format"""
    cleaned = {}

    # Clean date-based reservations
    for key, value in reservations.items():
        if key == '_client_memory':
            # Clean client memory
            cleaned_client_memory = {}
            for client_id, client_hours in value.items():
                cleaned_client_hours = {}
                for date_str, hour_str in client_hours.items():
                    # Clean Arabic format and generate proper format
                    if any(pattern in str(hour_str) for pattern in ['من', 'إلى', 'ىلإ', 'الى', ':00', 'h']):
                        # Extract hour numbers if possible, otherwise use 07-08 as default
                        import re
                        hour_match = re.search(r'(\d{2})', str(hour_str))
                        if hour_match:
                            hour_num = int(hour_match.group(1))
                            if 7 <= hour_num <= 16:
                                clean_hour = f"{hour_num:02d}-{hour_num+1:02d}"
                            else:
                                clean_hour = "07-08"  # Default fallback
                        else:
                            clean_hour = "07-08"  # Default fallback
                        cleaned_client_hours[date_str] = clean_hour
                        safe_print(f"[INFO] 🧹 Cleaned hour format: '{hour_str}' -> '{clean_hour}'")
                    else:
                        cleaned_client_hours[date_str] = hour_str
                cleaned_client_memory[client_id] = cleaned_client_hours
            cleaned[key] = cleaned_client_memory
        elif isinstance(value, list):
            # Clean hour lists for date reservations
            cleaned_hours = []
            for hour_str in value:
                if any(pattern in str(hour_str) for pattern in ['من', 'إلى', 'ىلإ', 'الى', ':00', 'h']):
                    # Extract hour numbers if possible
                    import re
                    hour_match = re.search(r'(\d{2})', str(hour_str))
                    if hour_match:
                        hour_num = int(hour_match.group(1))
                        if 7 <= hour_num <= 16:
                            clean_hour = f"{hour_num:02d}-{hour_num+1:02d}"
                        else:
                            clean_hour = "07-08"  # Default fallback
                    else:
                        clean_hour = "07-08"  # Default fallback
                    cleaned_hours.append(clean_hour)
                    safe_print(f"[INFO] 🧹 Cleaned reservation hour: '{hour_str}' -> '{clean_hour}'")
                else:
                    cleaned_hours.append(hour_str)
            cleaned[key] = cleaned_hours
        else:
            cleaned[key] = value

    return cleaned

def reserve_hours_for_dates(dates: List[str], client_id: str) -> List[str]:
    """Reserve unique hours per date within 07-17 window.

    Per-date logic with CLIENT MEMORY (CLEANED FORMAT):
    - Always generate clean XX-XX format (e.g., 07-08, 08-09)
    - Clean any old Arabic formats from client memory
    - First client on a date gets 07-08
    - Second client on same date gets 08-09
    - Third client on same date gets 09-10
    - etc. up to 16-17
    - Different dates reset hour assignment (start from 07-08 again)

    If all 10 slots are taken for that day, assign 'FULL' marker.
    """
    reservations = load_reservations()

    # Check if this client already has hour assignments
    client_memory_key = f'_client_memory'
    if client_memory_key not in reservations:
        reservations[client_memory_key] = {}

    client_memory = reservations[client_memory_key]

    # Clean up any old Arabic format entries for this client
    if client_id in client_memory:
        stored_hours = client_memory[client_id]
        cleaned_hours = {}

        for date_str, old_hour in stored_hours.items():
            # Check if hour is in old Arabic format and needs cleaning
            if any(pattern in str(old_hour) for pattern in ['من', 'إلى', 'ىلإ', 'الى', ':00', 'h']):
                # Generate new clean format for this date
                date_reservations = reservations.get(date_str, [])
                start_h = WORKING_HOURS['start']
                end_h = WORKING_HOURS['end']

                # Find first available slot
                assigned = None
                for hour in range(start_h, end_h):
                    slot = WORKING_HOURS['format_template'].format(start=hour, end=hour + 1)
                    if slot not in date_reservations:
                        date_reservations.append(slot)
                        assigned = slot
                        break

                if assigned is None:
                    assigned = 'FULL'

                cleaned_hours[date_str] = assigned
                reservations[date_str] = date_reservations
                safe_print(f"[INFO] 🧹 Cleaned old hour format for {client_id} on {date_str}: '{old_hour}' -> '{assigned}'")
            else:
                # Keep existing clean format
                cleaned_hours[date_str] = old_hour

        client_memory[client_id] = cleaned_hours

    # If client exists in memory, reuse their hours for matching dates
    if client_id in client_memory:
        stored_hours = client_memory[client_id]
        hours = []

        for date_str in dates:
            if date_str in stored_hours:
                # Ensure format is clean before reusing
                stored_hour = stored_hours[date_str]
                if any(pattern in str(stored_hour) for pattern in ['من', 'إلى', 'ىلإ', 'الى', ':00', 'h']):
                    # Regenerate clean format
                    date_reservations = reservations.get(date_str, [])
                    start_h = WORKING_HOURS['start']
                    end_h = WORKING_HOURS['end']

                    assigned = None
                    for hour in range(start_h, end_h):
                        slot = WORKING_HOURS['format_template'].format(start=hour, end=hour + 1)
                        if slot not in date_reservations:
                            date_reservations.append(slot)
                            assigned = slot
                            break

                    if assigned is None:
                        assigned = 'FULL'

                    hours.append(assigned)
                    stored_hours[date_str] = assigned
                    reservations[date_str] = date_reservations
                    safe_print(f"[INFO] 🧹 Regenerated clean hour for {client_id} on {date_str}: '{assigned}'")
                else:
                    # Reuse existing clean hour format
                    hours.append(stored_hour)
            else:
                # New date for existing client - assign next available hour
                date_reservations = reservations.get(date_str, [])
                assigned = None
                start_h = WORKING_HOURS['start']
                end_h = WORKING_HOURS['end']

                for hour in range(start_h, end_h):
                    slot = WORKING_HOURS['format_template'].format(start=hour, end=hour + 1)
                    if slot not in date_reservations:
                        date_reservations.append(slot)
                        assigned = slot
                        break

                if assigned is None:
                    assigned = 'FULL'

                hours.append(assigned)
                reservations[date_str] = date_reservations
                stored_hours[date_str] = assigned

        # Update memory
        client_memory[client_id] = stored_hours
        reservations[client_memory_key] = client_memory
        save_reservations(reservations)
        return hours

    # New client - assign hours and store in memory
    hours = []
    client_hours = {}
    start_h = WORKING_HOURS['start']
    end_h = WORKING_HOURS['end']

    for date_str in dates:
        date_reservations = reservations.get(date_str, [])
        assigned = None

        # Find first available hour slot for this specific date
        for hour in range(start_h, end_h):
            slot = WORKING_HOURS['format_template'].format(start=hour, end=hour + 1)
            if slot not in date_reservations:
                date_reservations.append(slot)
                assigned = slot
                break

        if assigned is None:
            assigned = 'FULL'

        hours.append(assigned)
        client_hours[date_str] = assigned
        reservations[date_str] = date_reservations

    # Store client memory
    client_memory[client_id] = client_hours
    reservations[client_memory_key] = client_memory
    save_reservations(reservations)
    return hours

def reserve_consistent_hour_for_dates(dates: List[str], client_id: str) -> List[str]:
    """Assign hours based on per-date availability (NEW LOGIC).

    Per-date assignment strategy:
      - For each date, find the next available hour slot (07-08, 08-09, 09-10, etc.)
      - First client on date X gets 07-08
      - Second client on same date X gets 08-09
      - Third client on same date X gets 09-10
      - Different date Y: Reset to 07-08 for first client on that date

    This ensures no two clients have the same hour on the same date,
    but allows reuse of hours across different dates.
    """
    # Use the per-date logic directly (no global assignments)
    return reserve_hours_for_dates(dates, client_id)

# Fixed placeholders for candidate follow-up card
PLACEHOLDERS = [
    'category', 'fullName', 'birthDate', 'birthPlace', 'address', 'phoneNumber', 'schoolSubmissionDate', 'vers'
]

# Mapping for misspelled placeholders in template
PLACEHOLDER_MAPPING = {
    'birtDate': 'birthDate',
    'birtPlace': 'birthPlace',
    'poneNumber': 'phoneNumber',
    'scoolSubmissionDate': 'schoolSubmissionDate'
}

WEEKEND_DAYS = {4, 5}  # الجمعة=4, السبت=5

# Tables configuration based on analysis
TABLES_CONFIG = {
    'personal_info': {'index': 0, 'has_placeholders': True},
    'exams': {'index': 1, 'skip': True},  # يبقى فارغ
    'lessons_table1': {'index': 2, 'rows': 30, 'dates_only': True, 'start_row': 2, 'date_col': 1},                    # Theory lessons - dates only - always filled
    'lessons_table2': {'index': 3, 'rows': 30, 'dates_and_hours': True, 'start_row': 2, 'date_col': 1, 'hour_col': 2}  # Practical lessons - dates + hours - only if traffic law passed
}

# Working hours configuration
WORKING_HOURS = {
    'start': 7,   # 07:00
    'end': 17,    # exclusive (last slot 16-17)
    'format_template': '{start:02d}-{end:02d}'
}

def parse_args():
    p = argparse.ArgumentParser(description='Fill Candidate Follow-up Card (Arabic).')
    p.add_argument('--input', default='resources/templates/بطاقة المتابعة للمترشح.docx', help='ملف Word المصدر.')
    p.add_argument('--output', help='ملف الإخراج (مطلوب).')
    p.add_argument('--start-date', dest='start_date', help='تاريخ بداية أول حصة (YYYY-MM-DD).')
    p.add_argument('--client-data', dest='client_data', help='بيانات العميل كـ JSON string.')
    p.add_argument('--data', dest='data_path', help='ملف JSON للبيانات العامة.')
    p.add_argument('--client-id', dest='client_id', help='معرف العميل الفريد لتجنب تعارض الساعات.')

    # Manual field overrides
    for ph in PLACEHOLDERS:
        p.add_argument(f'--{ph}', help=f'قيمة {ph} إذا لم يُستخدم --client-data')

    # Table configuration
    p.add_argument('--table1-dates', type=int, default=30, help='عدد التواريخ للجدول الأول (افتراضي 30).')
    p.add_argument('--table2-dates', type=int, default=30, help='عدد التواريخ للجدول الثاني (افتراضي 30).')
    p.add_argument('--date-format', default='%d/%m/%Y', help='تنسيق التاريخ (افتراضي %%d/%%m/%%Y).')

    # PDF options
    p.add_argument('--pdf', action='store_true', help='إنشاء نسخة PDF بالإضافة إلى DOCX.')
    p.add_argument('--pdf-only', action='store_true', help='إنشاء PDF فقط وحذف DOCX.')

    # Debug
    p.add_argument('--placeholders', action='store_true', help='اطبع قائمة الـ placeholders المتاحة ثم اخرج.')

    args = p.parse_args()

    # Validation
    if not args.placeholders:
        if not args.output:
            p.error('--output is required')
        if not args.start_date:
            p.error('--start-date is required')

    return args

def load_client_data(args) -> Tuple[Dict[str, str], bool]:
    """Load client data from various sources with proper priority. Returns (data, traffic_law_passed)"""
    data: Dict[str, str] = {}
    traffic_law_passed = False  # Default to False

    # 1. Load from JSON file if provided
    if args.data_path:
        try:
            with open(args.data_path, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
            if isinstance(loaded, dict):
                # Check if this looks like raw client data or processed placeholders
                if 'first_name_ar' in loaded or 'last_name_ar' in loaded:
                    # Raw client data format - convert to placeholders
                    client_data = {
                        'category': 'B',
                        'fullName': f"{loaded.get('first_name_ar', '')} {loaded.get('last_name_ar', '')}".strip(),
                        'birthDate': loaded.get('birth_date', ''),
                        'birthPlace': f"{loaded.get('birth_municipality', '')} {loaded.get('birth_state', '')}".strip(),
                        'address': ' '.join(filter(None, [
                            loaded.get('current_address', ''),
                            loaded.get('current_municipality', ''),
                            loaded.get('current_state', '')
                        ])),
                        'phoneNumber': loaded.get('phone_number', ''),
                        'schoolSubmissionDate': loaded.get('register_date', ''),
                        'vers': str(loaded.get('subPrice', '6000'))
                    }
                    data.update(client_data)

                    # Check traffic law test status from raw client data
                    tests = loaded.get('tests', {})
                    traffic_law_test = tests.get('trafficLawTest', {})
                    traffic_law_passed = traffic_law_test.get('passed', False)
                else:
                    # Already processed placeholder data
                    data.update({k: str(v) for k, v in loaded.items()})
        except Exception as e:
            print(f"[WARN] Failed to read JSON file: {e}")

    # 2. Load from client-data JSON string if provided
    if args.client_data:
        try:
            loaded = json.loads(args.client_data)
            if isinstance(loaded, dict):
                # INFO: Check traffic law test status for smart date calculation
                tests = loaded.get('tests', {})
                traffic_law_test = tests.get('trafficLawTest', {})
                traffic_law_passed = traffic_law_test.get('passed', False)

                if traffic_law_passed:
                    safe_print(f"[INFO] ✅ Traffic law test passed - using smart date calculation for: {loaded.get('first_name_ar', '')} {loaded.get('last_name_ar', '')}")
                else:
                    safe_print(f"[INFO] ⏳ Traffic law test not yet passed - generating template with empty practical lessons for: {loaded.get('first_name_ar', '')} {loaded.get('last_name_ar', '')}")

                # Convert client data format to placeholders format
                client_data = {
                    'category': 'B',  # Changed to English B as requested
                    'fullName': f"{loaded.get('first_name_ar', '')} {loaded.get('last_name_ar', '')}".strip(),
                    'birthDate': loaded.get('birth_date', ''),
                    'birthPlace': f"{loaded.get('birth_municipality', '')} {loaded.get('birth_state', '')}".strip(),
                    'address': ' '.join(filter(None, [
                        loaded.get('current_address', ''),
                        loaded.get('current_municipality', ''),
                        loaded.get('current_state', '')
                    ])),
                    'phoneNumber': loaded.get('phone_number', ''),
                    'schoolSubmissionDate': loaded.get('register_date', ''),
                    'vers': str(loaded.get('subPrice', '6000'))  # New placeholder for subPrice
                }
                data.update(client_data)
        except Exception as e:
            print(f"[WARN] Failed to read client-data JSON: {e}")

    # 3. Individual arguments have highest priority
    for ph in PLACEHOLDERS:
        val = getattr(args, ph, None)
        if val:
            data[ph] = val

    # 4. Default values as fallback
    defaults = {
        'category': 'B',  # Changed to English B
        'fullName': 'اسم المترشح',
        'birthDate': '1990/01/01',
        'birthPlace': 'المدينة',
        'address': 'العنوان',
        'phoneNumber': '0000000000',
        'schoolSubmissionDate': datetime.now().strftime('%Y/%m/%d'),
        'vers': '6000'  # Default subPrice
    }

    for k, v in defaults.items():
        data.setdefault(k, v)

    return data, traffic_law_passed

def generate_dates(start_date_str: str, count: int, out_format: str) -> List[str]:
    """Generate dates skipping weekends (same logic as traffic law lessons)"""
    # Try multiple date formats to be flexible
    date_formats = ['%Y-%m-%d', '%m/%d/%Y', '%d/%m/%Y', '%Y/%m/%d']
    current = None

    for fmt in date_formats:
        try:
            current = datetime.strptime(start_date_str, fmt).date()
            break
        except ValueError:
            continue

    if current is None:
        raise ValueError('تنسيق start-date يجب أن يكون YYYY-MM-DD أو MM/DD/YYYY أو DD/MM/YYYY')

    dates: List[str] = []
    while len(dates) < count:
        # Skip weekends (Friday=4, Saturday=5)
        if current.weekday() not in WEEKEND_DAYS:
            try:
                dates.append(current.strftime(out_format))
            except Exception:
                # Fallback safe format
                dates.append(current.strftime('%d/%m/%Y'))
        current += timedelta(days=1)
    return dates

def generate_hour_schedule(count: int, dates: List[str], client_id: str = None) -> List[str]:
    """Generate hour schedule within 07-17 window.

    With client_id -> ALL lessons get the SAME hour slot (consistent per client).
    Without client_id -> deterministic cycling inside window (no overflow past 16-17).
    """
    if client_id:
        # Get consistent hour assignment - SAME hour for ALL lessons
        consistent = reserve_consistent_hour_for_dates(dates[:count], client_id)
        return consistent

    # Fallback: without client_id, cycle hours per lesson
    hours: List[str] = []
    start_h = WORKING_HOURS['start']
    end_h = WORKING_HOURS['end']
    span = end_h - start_h
    for i in range(count):
        h = start_h + (i % span)
        hours.append(WORKING_HOURS['format_template'].format(start=h, end=h + 1))
    return hours

def replace_paragraph_placeholders(doc, data: Dict[str, str]):
    """Replace placeholders in paragraphs, tables, and XML elements (with mapping for misspelled placeholders)"""
    # Create extended data with misspelled placeholder mapping
    extended_data = data.copy()
    for misspelled, correct in PLACEHOLDER_MAPPING.items():
        if correct in data:
            extended_data[misspelled] = data[correct]

    # Replace in paragraphs
    for para in doc.paragraphs:
        if '{{' in para.text:
            original = para.text
            new_text = original

            # Use flexible regex to find placeholders even if malformed
            placeholders_in_text = re.findall(r'{{(\w+)\}?\}?', original)

            for key in placeholders_in_text:
                if key in extended_data:
                    value = extended_data.get(key, '')
                    # Try different placeholder patterns
                    patterns = [
                        '{{' + key + '}}',    # {{key}}
                        '{{' + key + '}',     # {{key}
                        '{{' + key + '}}}',   # {{key}}}
                    ]
                    replaced = False
                    for pattern in patterns:
                        if pattern in new_text:
                            new_text = new_text.replace(pattern, value)
                            replaced = True
                            break

                    if not replaced:
                        # Last attempt with regex
                        pattern = '{{' + key + r'\}?\}?'
                        new_text = re.sub(pattern, value, new_text)

            if new_text != original:
                para.clear()
                para.add_run(new_text)

    # Replace in tables (personal info table contains placeholders)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '{{' in para.text:
                        original = para.text
                        new_text = original

                        placeholders_in_text = re.findall(r'{{(\w+)\}?\}?', original)

                        for key in placeholders_in_text:
                            if key in extended_data:
                                value = extended_data.get(key, '')
                                patterns = [
                                    '{{' + key + '}}',
                                    '{{' + key + '}',
                                    '{{' + key + '}}}',
                                ]
                                replaced = False
                                for pattern in patterns:
                                    if pattern in new_text:
                                        new_text = new_text.replace(pattern, value)
                                        replaced = True
                                        break

                                if not replaced:
                                    pattern = '{{' + key + r'\}?\}?'
                                    new_text = re.sub(pattern, value, new_text)

                        if new_text != original:
                            para.clear()
                            para.add_run(new_text)

    # Replace hardcoded time format patterns throughout the document
    clean_hour_formats_in_document(doc)

    # Replace in XML elements using simple iteration
    try:
        # Find all text elements in the document
        for elem in doc._element.iter():
            if hasattr(elem, 'text') and elem.text and '{{' in elem.text:
                try:
                    original_text = elem.text
                    new_text = original_text

                    placeholders_in_text = re.findall(r'{{(\w+)\}?\}?', original_text)
                    for key in placeholders_in_text:
                        if key in extended_data:
                            value = extended_data.get(key, '')
                            patterns = [
                                '{{' + key + '}}',
                                '{{' + key + '}',
                                '{{' + key + '}}}',
                            ]
                            for pattern in patterns:
                                if pattern in new_text:
                                    new_text = new_text.replace(pattern, value)
                                    break

                    if new_text != original_text:
                        # Try to set the text, skip if readonly
                        elem.text = new_text
                        safe_print(f"[INFO] XML replacement: '{original_text}' -> '{new_text}'")
                except:
                    # Skip readonly elements
                    pass

    except Exception as e:
        print(f"[WARN] Failed to process XML elements: {e}")

def clean_hour_formats_in_document(doc):
    """Remove hardcoded hour format patterns from the document - SMART CLEANING

    Only clean hardcoded Arabic patterns from template, but preserve cells that will be filled with proper hours
    """
    # Enhanced patterns to catch all possible Arabic time formats
    hour_patterns = [
        r'من\s*\d{2}:\d{2}\s*إلى\s*\d{2}:\d{2}h?',  # من07:00إلى08:00h
        r'نم\s*\d{2}:\d{2}\s*ىلإ\s*\d{2}:\d{2}h?',  # نم07:00ىلإ08:00h
        r'15الى16سا',                                # Hardcoded 15الى16سا pattern
        r'\d+الى\d+سا',                             # Any number الى number سا pattern
    ]

    total_cleaned = 0

    # Focus specifically on Table 2 (theory lessons) and Table 3 (practical lessons)
    # where the hour cells are located - but only clean hardcoded Arabic patterns
    if len(doc.tables) > 2:
        for table_idx in [2, 3]:  # Table 2 and Table 3
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                safe_print(f"[INFO] 🧹 Smart cleaning of hardcoded Arabic patterns in Table {table_idx}")

                for row_idx, row in enumerate(table.rows):
                    if row_idx >= 2:  # Skip header rows
                        for cell_idx, cell in enumerate(row.cells):
                            if cell_idx == 2:  # Hour column (index 2)
                                original = cell.text.strip()
                                if original:
                                    # Only clean if it matches hardcoded Arabic patterns
                                    has_arabic_pattern = any(re.search(pattern, original) for pattern in hour_patterns)
                                    if has_arabic_pattern:
                                        # Completely clear the cell content and all formatting
                                        cell._element.clear_content()
                                        total_cleaned += 1
                                        safe_print(f"[INFO] 🧹 CLEARED hardcoded pattern in Table {table_idx}, Row {row_idx}, Cell {cell_idx}: '{original}' -> EMPTY")

    # Also clean header cells that might have Arabic text
    # Clean "الساعة" (hour) to "العة" as seen in debug output
    arabic_header_patterns = [
        (r'الساعة', 'العة'),  # Clean Arabic "hour" header
    ]

    # Clean XML elements - more targeted approach for headers
    try:
        for elem in doc._element.iter():
            if hasattr(elem, 'text') and elem.text:
                original_text = elem.text
                new_text = original_text

                for pattern, replacement in arabic_header_patterns:
                    new_text = re.sub(pattern, replacement, new_text)

                if new_text != original_text:
                    try:
                        elem.text = new_text
                        total_cleaned += 1
                        safe_print(f"[INFO] 🧹 Cleaned XML header: '{original_text}' -> '{new_text}'")
                    except:
                        pass  # Skip readonly elements
    except Exception as e:
        print(f"[WARN] Failed to clean hour formats in XML elements: {e}")

    if total_cleaned > 0:
        safe_print(f"[INFO] ✅ Total hardcoded patterns cleaned: {total_cleaned}")
    else:
        safe_print(f"[INFO] ✅ No hardcoded patterns found to clean")

def clean_text(text: str) -> str:
    """Clean text by removing extra spaces, newlines, and trimming"""
    if not text:
        return ""
    return re.sub(r'\s+', ' ', text.strip())

def format_cell(cell, value: str):
    """Apply compact formatting: Arial size 12, bold, center aligned, completely replace cell content."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    cleaned = clean_text(value)

    # COMPLETELY clear the cell - remove all paragraphs
    cell._element.clear_content()

    # Add a fresh paragraph
    para = cell.add_paragraph()
    run = para.add_run(cleaned)

    # Set paragraph alignment
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set paragraph formatting
    if para.paragraph_format is not None:
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1

    # Set run formatting
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True

    # Set cell vertical alignment
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        # Remove existing margins
        for child in list(tcPr):
            if child.tag.endswith('tcMar'):
                tcPr.remove(child)
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'start', 'bottom', 'end']:
            side_el = OxmlElement(f'w:{side}')
            side_el.set(qn('w:w'), '0')
            side_el.set(qn('w:type'), 'dxa')
            tcMar.append(side_el)
        tcPr.append(tcMar)
    except Exception:
        pass

    return cleaned

def fill_table_dates_only(table, dates: List[str], start_row: int, date_col: int):
    """Fill table with dates only (for 30 lessons table)"""
    filled_count = 0
    for i, date_str in enumerate(dates):
        row_idx = start_row + i
        if row_idx >= len(table.rows):
            break

        row = table.rows[row_idx]
        if date_col >= len(row.cells):
            break

        # Compact format date cell
        date_cell = row.cells[date_col]
        format_cell(date_cell, date_str)

        filled_count += 1

    return filled_count

def fill_table_dates_and_hours(table, dates: List[str], hours: List[str], start_row: int, date_col: int, hour_col: int):
    """Fill table with dates and hours (for 30 lessons table)"""
    filled_count = 0
    for i in range(min(len(dates), len(hours))):
        row_idx = start_row + i
        if row_idx >= len(table.rows):
            break

        row = table.rows[row_idx]
        if date_col >= len(row.cells) or hour_col >= len(row.cells):
            break

        # Compact format date and hour cells
        date_cell = row.cells[date_col]
        format_cell(date_cell, dates[i])

        hour_cell = row.cells[hour_col]
        format_cell(hour_cell, hours[i])

        filled_count += 1

    return filled_count

def fill_candidate_tables(doc, start_date_str: str, table1_count: int, table2_count: int, date_format: str, client_id: str = None, traffic_law_passed: bool = False):
    """Fill the candidate follow-up card tables based on discovered structure

    Table 2 (Theory lessons): Always filled with DATES ONLY (no hours)
    Table 3 (Practical lessons): Only filled with DATES + HOURS if traffic law test passed
    """
    results = {}

    # Generate dates for table 1 (theory lessons - always filled, dates only)
    table1_dates = generate_dates(start_date_str, table1_count, date_format)

    # Fill Table 2 (lessons_table1) - theory lessons with DATES ONLY (always filled)
    config_table1 = TABLES_CONFIG['lessons_table1']
    if config_table1['index'] < len(doc.tables):
        table1 = doc.tables[config_table1['index']]
        filled_table1 = fill_table_dates_only(
            table1, table1_dates,
            config_table1['start_row'], config_table1['date_col']
        )
        results['table1'] = f"Filled {filled_table1} theory lesson date rows of {table1_count}"
        safe_print(f"[INFO] ✅ Theory lessons table filled with dates only")
    else:
        results['table1'] = f"Table 2 not found in document"

    # Fill Table 3 (lessons_table2) - practical lessons with dates and hours
    # Only fill if traffic law test has been passed
    config_table2 = TABLES_CONFIG['lessons_table2']
    if traffic_law_passed and config_table2['index'] < len(doc.tables):
        # Calculate start date for table 2: 7 days after the last date in table1
        # Parse the last date using the same flexible parsing as generate_dates
        last_table1_date_str = table1_dates[-1]
        date_formats_to_try = ['%Y-%m-%d', '%m/%d/%Y', '%d/%m/%Y', '%Y/%m/%d']
        last_table1_date = None

        for fmt in date_formats_to_try:
            try:
                last_table1_date = datetime.strptime(last_table1_date_str, fmt)
                break
            except ValueError:
                continue

        if last_table1_date is None:
            # Fallback: assume it's in the same format we're using for output
            last_table1_date = datetime.strptime(last_table1_date_str, date_format)

        table2_start_date = last_table1_date + timedelta(days=7)
        table2_start_str = table2_start_date.strftime('%Y-%m-%d')  # Always use standard format for input

        # Generate dates for table 2 (starting 7 days after table1 ends)
        table2_dates = generate_dates(table2_start_str, table2_count, date_format)

        # Generate smart hours for table 2 using reservation system
        table2_hours = generate_hour_schedule(table2_count, table2_dates, client_id)

        table2 = doc.tables[config_table2['index']]
        filled_table2 = fill_table_dates_and_hours(
            table2, table2_dates, table2_hours,
            config_table2['start_row'], config_table2['date_col'], config_table2['hour_col']
        )
        results['table2'] = f"Filled {filled_table2} practical lesson date/hour rows of {table2_count}"
        safe_print(f"[INFO] ✅ Practical lessons table filled with dates and hours - student has passed traffic law test")
    else:
        if not traffic_law_passed:
            safe_print(f"[INFO] ⏳ Practical lessons table left empty - student has not yet passed traffic law test")
            results['table2'] = f"Skipped practical lessons table - traffic law test not passed"
        else:
            safe_print(f"[WARN] Table 3 index {config_table2['index']} not found in document")
            results['table2'] = f"Table 3 not found in document"

    return results

def main():
    args = parse_args()

    if args.placeholders:
        print('\n'.join(PLACEHOLDERS))
        return

    # Load client data and traffic law test status
    data, traffic_law_passed = load_client_data(args)
    safe_print(f"[INFO] Data loaded with keys: {list(data.keys())}")

    # Validate input file
    input_path = Path(args.input)
    if not input_path.exists():
        print(f"[ERROR] Source file not found: {input_path}", file=sys.stderr)
        sys.exit(1)

    # Load document
    doc = Document(str(input_path))
    print(f"[INFO] Document loaded: {len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs")

    # Replace placeholders
    replace_paragraph_placeholders(doc, data)
    print("[INFO] Placeholders replaced successfully")

    # Fill tables with dates and hours
    results = fill_candidate_tables(
        doc, args.start_date,
        args.table1_dates, args.table2_dates,
        args.date_format, args.client_id,  # Pass client ID
        traffic_law_passed  # Pass traffic law test status
    )

    for table_name, result in results.items():
        safe_print(f"[INFO] {table_name}: {result}")

    # Apply Arial font size 8 bold to table content only
    try:
        from docx.shared import Pt
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Arial'
                            run.font.bold = True
                            run.font.size = Pt(10)
        print("[INFO] Applied Arial font size 8 bold to table content")
    except Exception as e:
        print(f"[WARN] Failed to apply font formatting: {e}")
    # Save output with better Unicode path handling
    output_path = Path(args.output)
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Try to save with unicode handling
        output_str = str(output_path)

        # If path contains non-ASCII characters, try to create a safe filename
        if not output_str.isascii():
            # Create a timestamp-based safe filename in the same directory
            import time
            timestamp = int(time.time())
            safe_name = f"candidate_follow_up_{timestamp}.docx"
            safe_output_path = output_path.parent / safe_name
            print(f"[WARN] Using safe filename due to Unicode path: {safe_output_path}")
            doc.save(str(safe_output_path))
            safe_print(f"[OK] File created: {safe_output_path}")

            # Try to rename to original if possible
            try:
                safe_output_path.rename(output_path)
                safe_print(f"[OK] Renamed to original: {output_path}")
            except Exception as rename_err:
                print(f"[WARN] Could not rename to original path: {rename_err}")
                safe_print(f"[OK] Final file: {safe_output_path}")
        else:
            doc.save(output_str)
            safe_print(f"[OK] File created: {output_path}")

    except PermissionError as e:
        print(f"[ERROR] Permission denied. File may be open in another application: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print(f"[ERROR] Failed to save file: {e}", file=sys.stderr)
        sys.exit(1)

    # Convert to PDF if requested
    if args.pdf or args.pdf_only:
        if convert is None:
            print("[ERROR] Need to install docx2pdf library: pip install docx2pdf", file=sys.stderr)
            sys.exit(1)

        pdf_path = output_path.with_suffix('.pdf')
        try:
            # Ensure the DOCX file exists before converting
            if not output_path.exists():
                print(f"[ERROR] DOCX file not found for conversion: {output_path}", file=sys.stderr)
                sys.exit(1)

            # Ensure output directory exists for PDF
            pdf_path.parent.mkdir(parents=True, exist_ok=True)

            convert(str(output_path), str(pdf_path))
            safe_print(f"[OK] PDF file created: {pdf_path}")
            print(f"PDF_PATH={pdf_path}")

            # Delete DOCX if pdf-only requested
            if args.pdf_only:
                output_path.unlink(missing_ok=True)
                safe_print(f"[OK] DOCX file deleted (pdf-only mode)")
        except Exception as e:
            print(f"[ERROR] Failed to convert to PDF: {e}", file=sys.stderr)
            # Don't exit on PDF conversion failure, return DOCX instead
            print(f"[WARN] Will use DOCX file instead of PDF")
            print(f"PDF_PATH={output_path}")  # Return DOCX path as fallback

if __name__ == '__main__':
    main()
