# -*- coding: utf-8 -*-
"""
ملء استمارة "بطاقة خاصة بدروس قانون المرور" باستبدال الحقول النصية العلوية
وإدراج تواريخ 30 حصة في العمود الثالث من أول جدول.

المتطلبات (وفق الوصف):
- المستند المصدر يحتوي على فقرات في الأعلى تضم Placeholders:
    {{fullName}}, {{birthDate}}, {{birthPlace}}, {{address}}, {{registrationDate}}
- أول جدول (doc.tables[0]) يحوي 30 صف مخصصة للدروس (نفترض أن الصفوف المطلوب تعبئتها هي كلها أو نقدر العدد الفعلي)
- نملأ العمود الثالث (index=2) بالتواريخ المتسلسلة انطلاقاً من تاريخ بداية (start-date)
- تخطي (عدم جدولة) يومي الجمعة والسبت.
- إخراج ملف جديد: بطاقة_قانون_المرور_مملوءة.pdf (PDF افتراضي)

الاستخدام (مثال):
python scripts/fill_traffic_law_lessons_card.py \
  --input "resources/templates/بطاقة خاصة بدروس قانون المرور .docx" \
  --output "output/بطاقة_قانون_المرور_مملوءة.pdf" \
  --start-date 2025-01-12 \
  --data data/traffic_lessons_card.json

مثال JSON:
{
  "fullName": "أحمد بن صالح",
  "birthDate": "2000/05/12",
  "birthPlace": "المسيلة",
  "address": "حي النصر، المسيلة",
  "registrationDate": "2025/01/10"
}

ملاحظات:
- الآن يتم إنشاء ملف PDF افتراضياً بدلاً من DOCX
- في حالة عدم تقديم ملف JSON يمكن تمرير القيم عبر وسيطات منفصلة أو استخدام قيم افتراضية.
- إذا كان الجدول يحتوي صف رأس (Header) يمكن استخدام العلم --has-header لتخطي أول صف.
- استخدم --docx-only لإنشاء DOCX فقط أو --docx للاحتفاظ بكل من PDF وDOCX.
- السكربت لا يعدل Placeholder داخل الجدول (غير مطلوبة).
"""
from __future__ import annotations
import argparse
import json
import sys
import re
from pathlib import Path
from datetime import datetime, timedelta
from typing import Dict, Any, List

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:  # pragma: no cover
    print("[ERROR] تحتاج لتثبيت المكتبة python-docx: pip install python-docx", file=sys.stderr)
    sys.exit(1)

try:
    from docx2pdf import convert
except ImportError:  # pragma: no cover
    convert = None

PLACEHOLDERS = [
    'fullName', 'birthDate', 'birthPlace', 'address', 'registrationDate'
]

WEEKEND_DAYS = [4, 5]  # الجمعة=4, السبت=5


def clean_text(text):
    """Remove extra spaces and newlines from text."""
    if not text:
        return text
    return re.sub(r'\s+', ' ', str(text)).strip()

def format_cell(cell, value: str):
    """Apply strict compact formatting: single run, Arial 7, no extra spacing, removed cell margins."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    cleaned = clean_text(value)
    # Remove extra paragraphs
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._p
        p.getparent().remove(p)
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(cleaned)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if para.paragraph_format is not None:
        pf = para.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    run.font.bold = True
    try:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        # Remove existing tcMar
        for child in list(tcPr):
            if child.tag.endswith('tcMar'):
                tcPr.remove(child)
        tcMar = OxmlElement('w:tcMar')
        for side in ['top', 'start', 'bottom', 'end']:
            side_el = OxmlElement(f'w:{side}')
            side_el.set(qn('w:w'), '0')
            side_el.set(qn('w:type'), 'dxa')
            tcMar.append(side_el)
        tcPr.append(tcMar)
    except Exception:
        pass
    return cleaned

def generate_dates(start_date_str: str, count: int, out_format: str) -> List[str]:
    """Generate sequential dates skipping Friday/Saturday."""
    try:
        # Handle both YYYY-MM-DD and YYYY/MM/DD formats
        if '/' in start_date_str:
            current = datetime.strptime(start_date_str, '%Y/%m/%d').date()
        else:
            current = datetime.strptime(start_date_str, '%Y-%m-%d').date()
    except ValueError:
        raise ValueError('تنسيق start-date يجب أن يكون YYYY-MM-DD أو YYYY/MM/DD')

    dates: List[str] = []
    while len(dates) < count:
        if current.weekday() not in WEEKEND_DAYS:
            try:
                dates.append(current.strftime(out_format))
            except Exception:
                dates.append(current.strftime('%Y/%m/%d'))
        current += timedelta(days=1)
    return dates

def parse_args():
    p = argparse.ArgumentParser(description='Fill Traffic Law Lessons Card (Arabic).')
    p.add_argument('--input', help='ملف Word المصدر.')
    p.add_argument('--output', default='output/بطاقة_قانون_المرور_مملوءة.pdf', help='ملف الإخراج.')
    p.add_argument('--start-date', dest='start_date', help='تاريخ بداية أول حصة (YYYY-MM-DD أو YYYY/MM/DD).')
    p.add_argument('--data', dest='data_path', help='ملف JSON للبيانات العامة (حقول الرأس).')
    for ph in PLACEHOLDERS:
        p.add_argument(f'--{ph}', help=f'قيمة {ph} إذا لم يُستخدم --data')
    p.add_argument('--sessions', type=int, default=30, help='عدد الحصص المطلوب جدولتها (افتراضي 30).')
    p.add_argument('--has-header', action='store_true', help='إذا كان أول صف في الجدول عبارة عن رأس يتم تخطيه.')
    p.add_argument('--table-index', type=int, default=0, help='رقم الجدول (index) الذي يحتوي على مواعيد الحصص (افتراضي 0).')
    p.add_argument('--date-column', type=int, default=2, help='رقم العمود (index) الذي تُكتب فيه تواريخ الحصص (افتراضي 2 للعمود الثالث).')
    p.add_argument('--date-format', default='%Y/%m/%d', help='تنسيق التاريخ داخل الجدول (افتراضي %%Y/%%m/%%d).')
    p.add_argument('--placeholders', action='store_true', help='اطبع قائمة الـ placeholders المتاحة ثم اخرج.')
    p.add_argument('--docx', action='store_true', help='إنشاء نسخة DOCX بالإضافة إلى PDF.')
    p.add_argument('--docx-only', action='store_true', help='إنشاء DOCX فقط وتخطي PDF.')

    args = p.parse_args()
    if not args.placeholders:
        if not args.input:
            p.error('--input is required')
        if not args.start_date:
            p.error('--start-date is required')
    return args


def load_data(args) -> Dict[str, str]:
    data: Dict[str, str] = {}
    if args.data_path:
        try:
            with open(args.data_path, 'r', encoding='utf-8') as f:
                loaded = json.load(f)
            if not isinstance(loaded, dict):
                raise ValueError('JSON يجب أن يكون كائن (object).')
            data.update({k: str(v) for k, v in loaded.items()})
        except Exception as e:
            print(f"[WARN] تعذر قراءة JSON: {e}")

    # وسيطات فردية لها أولوية
    for ph in PLACEHOLDERS:
        val = getattr(args, ph, None)
        if val:
            data[ph] = val

    # قيم افتراضية احتياطية في حال غياب بعضها - فقط للمفاتيح الفارغة
    defaults = {
        'fullName': 'اسم المتدرب',
        'birthDate': '1990/01/01',
        'birthPlace': 'المدينة',
        'address': 'العنوان',
        'registrationDate': '2025/01/01'
    }
    for k, v in defaults.items():
        if not data.get(k):  # Only use default if key is missing or empty
            data.setdefault(k, v)

    return data



def replace_paragraph_placeholders(doc, data: Dict[str, str]):
    # استبدال في الفقرات العادية
    for para in doc.paragraphs:
        if '{{' in para.text:
            original = para.text
            new_text = original

            # استخدام regex مرن للعثور على placeholders حتى لو كانت معطوبة
            placeholders_in_text = re.findall(r'{{(\w+)\}?\}?', original)

            # استبدال كل placeholder موجود
            for key in placeholders_in_text:
                if key in data:
                    value = data.get(key, '')
                    # جرب الأشكال المختلفة للـ placeholder
                    patterns = [
                        '{{' + key + '}}',    # {{key}}
                        '{{' + key + '}',     # {{key}
                        '{{' + key + '}}}',   # {{key}}}
                    ]
                    replaced = False
                    for pattern in patterns:
                        if pattern in new_text:
                            new_text = new_text.replace(pattern, value)
                            replaced = True
                            break

                    if not replaced:
                        # محاولة أخيرة: استبدال باستخدام regex
                        pattern = '{{' + key + r'\}?\}?'
                        new_text = re.sub(pattern, value, new_text)

            if new_text != original:
                para.clear()
                para.add_run(new_text)

    # استبدال في الجداول (نفس المنطق)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if '{{' in para.text:
                        original = para.text
                        new_text = original

                        placeholders_in_text = re.findall(r'{{(\w+)\}?\}?', original)

                        for key in placeholders_in_text:
                            if key in data:
                                value = data.get(key, '')
                                patterns = [
                                    '{{' + key + '}}',
                                    '{{' + key + '}',
                                    '{{' + key + '}}}',
                                ]
                                replaced = False
                                for pattern in patterns:
                                    if pattern in new_text:
                                        new_text = new_text.replace(pattern, value)
                                        replaced = True
                                        break

                                if not replaced:
                                    pattern = '{{' + key + r'\}?\}?'
                                    new_text = re.sub(pattern, value, new_text)

                        if new_text != original:
                            para.clear()
                            para.add_run(new_text)


def fill_table_dates(doc, dates: List[str], has_header: bool, table_index: int, date_column: int):
    if not doc.tables:
        raise ValueError('لم يتم العثور على أي جدول في المستند.')
    if table_index < 0 or table_index >= len(doc.tables):
        raise ValueError(f'رقم الجدول خارج النطاق. المستند يحتوي على {len(doc.tables)} جداول.')
    table = doc.tables[table_index]
    start_index = 1 if has_header else 0

    dates_filled = 0
    for i, date_str in enumerate(dates):
        # تأكد من وجود صف كافٍ
        if (start_index + i) >= len(table.rows):
            break

        row = table.rows[start_index + i]
        if date_column < 0 or date_column >= len(row.cells):
            raise ValueError(f'الصف {start_index + i} لا يحتوي على عمود رقم {date_column}. عدد الأعمدة المتاحة: {len(row.cells)}')

        # Compact date cell
        format_cell(row.cells[date_column], date_str)
        # Try remove fixed row height
        try:
            from docx.enum.table import WD_ROW_HEIGHT_RULE
            row.height = None
            row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
        except Exception:
            pass
        dates_filled += 1

    print(f"[INFO] Filled {dates_filled} date rows with Arial font size 7")


def main():
    args = parse_args()
    if args.placeholders:
        print('\n'.join(PLACEHOLDERS))
        return

    data = load_data(args)
    dates = generate_dates(args.start_date, args.sessions, args.date_format)

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"[ERROR] الملف المصدر غير موجود: {input_path}", file=sys.stderr)
        sys.exit(1)

    doc = Document(str(input_path))
    replace_paragraph_placeholders(doc, data)
    fill_table_dates(doc, dates, args.has_header, args.table_index, args.date_column)

    # Create temporary DOCX file
    output_path = Path(args.output)
    if output_path.suffix.lower() == '.pdf':
        docx_path = output_path.with_suffix('.docx')
    else:
        docx_path = output_path

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))

    # Determine what output to generate
    if args.docx_only:
        # DOCX only
        print(f"[OK] تم إنشاء الملف: {docx_path}")
        print(f"DOCX_PATH={docx_path}")
    else:
        # PDF by default (unless --docx-only specified)
        if convert is None:
            print("[ERROR] تحتاج لتثبيت المكتبة docx2pdf: pip install docx2pdf", file=sys.stderr)
            sys.exit(1)

        pdf_path = output_path if output_path.suffix.lower() == '.pdf' else output_path.with_suffix('.pdf')
        try:
            convert(str(docx_path), str(pdf_path))
            print(f"[OK] تم إنشاء ملف PDF: {pdf_path}")
            print(f"PDF_PATH={pdf_path}")

            # Keep DOCX if --docx flag specified, otherwise delete it
            if args.docx:
                print(f"[OK] تم الاحتفاظ بملف DOCX: {docx_path}")
                print(f"DOCX_PATH={docx_path}")
            else:
                docx_path.unlink(missing_ok=True)
                print(f"[OK] تم حذف ملف DOCX المؤقت")
        except Exception as e:
            print(f"[ERROR] فشل في تحويل PDF: {e}", file=sys.stderr)
            # Fallback to DOCX if PDF conversion fails
            print(f"[FALLBACK] سيتم استخدام ملف DOCX: {docx_path}")
            print(f"DOCX_PATH={docx_path}")
            return


if __name__ == '__main__':  # pragma: no cover
    main()
