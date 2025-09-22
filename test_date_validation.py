#!/usr/bin/env python3
"""
Test script to validate date formats in generated Word documents
"""
import sys
from pathlib import Path
sys.path.append(str(Path(__file__).parent))

try:
    from docx import Document
    import re
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

def check_date_formats(docx_path):
    """Check if all dates in the document follow dd/mm/yyyy format"""

    print(f"🔍 Checking date formats in: {docx_path}")

    doc = Document(docx_path)
    date_issues = []
    dd_mm_yyyy_count = 0
    other_format_count = 0

    # Date pattern for dd/mm/yyyy (01/01/2025, 31/12/2024, etc.)
    dd_mm_yyyy_pattern = r'\b(\d{1,2})/(\d{1,2})/(\d{4})\b'

    # Other date patterns to catch
    yyyy_mm_dd_pattern = r'\b(\d{4})-(\d{1,2})-(\d{1,2})\b'
    yyyy_slash_pattern = r'\b(\d{4})/(\d{1,2})/(\d{1,2})\b'

    def check_text_for_dates(text, location):
        nonlocal dd_mm_yyyy_count, other_format_count

        # Find dd/mm/yyyy format (correct)
        dd_mm_matches = re.findall(dd_mm_yyyy_pattern, text)
        for match in dd_mm_matches:
            day, month, year = match
            if 1 <= int(day) <= 31 and 1 <= int(month) <= 12:
                dd_mm_yyyy_count += 1
                print(f"  ✅ DD/MM/YYYY: {day}/{month}/{year} in {location}")

        # Find yyyy-mm-dd format (wrong for display)
        yyyy_dash_matches = re.findall(yyyy_mm_dd_pattern, text)
        for match in yyyy_dash_matches:
            year, month, day = match
            other_format_count += 1
            date_issues.append(f"  ❌ YYYY-MM-DD: {year}-{month}-{day} in {location} (should be DD/MM/YYYY)")

        # Find yyyy/mm/dd format (wrong for display)
        yyyy_slash_matches = re.findall(yyyy_slash_pattern, text)
        for match in yyyy_slash_matches:
            year, month, day = match
            other_format_count += 1
            date_issues.append(f"  ❌ YYYY/MM/DD: {year}/{month}/{day} in {location} (should be DD/MM/YYYY)")

    # Check paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            check_text_for_dates(para.text, f"paragraph {i+1}")

    # Check tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    if para.text.strip():
                        check_text_for_dates(
                            para.text,
                            f"table {table_idx+1}, row {row_idx+1}, cell {cell_idx+1}"
                        )

    # Summary
    print(f"\n📊 Date Format Summary:")
    print(f"  ✅ DD/MM/YYYY format: {dd_mm_yyyy_count} dates")
    print(f"  ❌ Other formats: {other_format_count} dates")

    if date_issues:
        print(f"\n🚨 Issues found:")
        for issue in date_issues:
            print(issue)
        return False
    else:
        print(f"\n🎉 All dates are in correct DD/MM/YYYY format!")
        return True

if __name__ == "__main__":
    docx_file = "test_final_date_check.docx"

    if not Path(docx_file).exists():
        print(f"ERROR: File {docx_file} not found")
        sys.exit(1)

    success = check_date_formats(docx_file)
    sys.exit(0 if success else 1)
